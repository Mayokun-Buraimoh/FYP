"""
Convert manuscript HTML (from Document.manuscript_content) to a Word .docx file.
Preserves paragraph flow, indentation, bold, headings, and page breaks from the PDF.
"""

import re
from html import escape
from io import BytesIO

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document as DocxDocument
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt

from .manuscript_html import normalize_manuscript_html
from .reference_list import strip_references_section_from_html

_BLOCK_TAGS = frozenset({"p", "h1", "h2", "h3", "h4", "h5", "h6", "div"})
_HEADING_TAGS = frozenset({"h1", "h2", "h3", "h4", "h5", "h6"})


def _paragraphs_from_manuscript_content(manuscript_content):
    if not manuscript_content:
        return []
    if isinstance(manuscript_content, list):
        return manuscript_content
    return []


def _manuscript_html(manuscript_content):
    paragraphs = _paragraphs_from_manuscript_content(manuscript_content)
    html_parts = []
    text_parts = []
    for block in paragraphs:
        if not isinstance(block, dict):
            continue
        html = (block.get("html") or "").strip()
        text = (block.get("text") or "").strip()
        if html:
            html_parts.append(html)
        elif text:
            text_parts.append(text)

    if html_parts:
        raw = "\n".join(html_parts)
    elif text_parts:
        plain = "\n\n".join(text_parts)
        raw = "\n".join(f"<p>{escape(p)}</p>" for p in plain.split("\n\n") if p.strip())
    else:
        return ""

    return normalize_manuscript_html(raw)


def _parse_margin_left_px(style_value):
    if not style_value:
        return 0
    match = re.search(r"margin-left:\s*([\d.]+)px", style_value, re.I)
    if not match:
        return 0
    try:
        return float(match.group(1))
    except ValueError:
        return 0


def _configure_document_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)


def _add_runs_to_paragraph(paragraph, node, bold=False, italic=False):
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
        return

    if not isinstance(node, Tag):
        return

    tag = node.name.lower() if node.name else ""
    if tag == "br":
        paragraph.add_run().add_break()
        return

    child_bold = bold or tag in ("strong", "b")
    child_italic = italic or tag in ("em", "i")

    for child in node.children:
        _add_runs_to_paragraph(paragraph, child, child_bold, child_italic)


def _apply_block_format(paragraph, tag, indent_px):
    fmt = paragraph.paragraph_format
    if indent_px > 0:
        fmt.left_indent = Inches(min(indent_px / 96.0, 2.25))

    if tag in _HEADING_TAGS:
        fmt.space_before = Pt(14)
        fmt.space_after = Pt(6)
        size = Pt(13 if tag == "h3" else 14)
        for run in paragraph.runs:
            run.bold = True
            run.font.size = size
    else:
        fmt.space_after = Pt(6)


def _add_block_element(doc, element):
    tag = (element.name or "p").lower()
    if tag == "div":
        return

    style_attr = element.get("style") or ""
    indent_px = _parse_margin_left_px(style_attr)
    p = doc.add_paragraph()
    _add_runs_to_paragraph(p, element)
    if not p.text.strip():
        return
    _apply_block_format(p, tag, indent_px)


def _iter_flow_blocks(root):
    for child in root.children:
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                yield ("p", "", text)
            continue
        if not isinstance(child, Tag):
            continue

        name = (child.name or "").lower()
        if child.get("data-page-break"):
            yield ("page-break", "", "")
            continue

        if name in _BLOCK_TAGS:
            if name == "div" and child.find(list(_BLOCK_TAGS)):
                yield from _iter_flow_blocks(child)
            elif name == "div":
                inner = child.get_text(separator=" ", strip=True)
                if inner:
                    yield ("p", child.get("style") or "", inner)
            else:
                yield (name, child.get("style") or "", child)
            continue


def _html_to_docx(doc, html_content):
    soup = BeautifulSoup(f"<div>{html_content}</div>", "html.parser")
    root = soup.find("div")
    if not root:
        return

    blocks = list(_iter_flow_blocks(root))
    if not blocks:
        text = root.get_text(separator="\n", strip=True)
        if text:
            doc.add_paragraph(text)
        return

    for kind, style, payload in blocks:
        if kind == "page-break":
            doc.add_page_break()
            continue

        if isinstance(payload, str):
            p = doc.add_paragraph(payload)
            _apply_block_format(p, "p", _parse_margin_left_px(style))
            continue

        _add_block_element(doc, payload)


def _append_references_section(doc, references):
    if not references:
        return
    doc.add_page_break()
    heading = doc.add_paragraph()
    heading_run = heading.add_run("References")
    heading_run.bold = True
    heading_run.font.size = Pt(14)
    heading.paragraph_format.space_after = Pt(12)

    for entry in references:
        p = doc.add_paragraph(entry.strip())
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(6)


def build_manuscript_docx(
    manuscript_content,
    title="Manuscript",
    references=None,
    strip_in_body_references=False,
):
    html_content = _manuscript_html(manuscript_content)
    if strip_in_body_references and html_content.strip():
        html_content = strip_references_section_from_html(html_content)
    doc = DocxDocument()
    _configure_document_styles(doc)

    if html_content.strip():
        _html_to_docx(doc, html_content)
    else:
        doc.add_paragraph("(No manuscript content yet.)")

    _append_references_section(doc, references or [])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
